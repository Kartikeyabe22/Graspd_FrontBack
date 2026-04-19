import os
import re

# -------------------- TEACHING JSON OBJECT EXTRACTION --------------------
# -------------------- SAFE JSON EXTRACTION --------------------

def extract_json_object(text: str):
    import json

    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1:
        return None

    json_str = text[start:end+1]

    try:
        return json.loads(json_str)
    except Exception:
        return None


def extract_json_array(text: str):
    import re, json

    match = re.search(r"\[.*?\]", text, re.DOTALL)
    if not match:
        return None

    try:
        return json.loads(match.group())
    except Exception:
        return None
# -------------------- DIRECT LLM (NO RAG) --------------------
# -------------------- DIRECT LLM --------------------

def run_llm_direct(prompt_text: str) -> str:
    if not api_key:
        raise Exception("Missing GROQ_API_KEY")

    llm = ChatGroq(
        groq_api_key=api_key,
        model_name="llama-3.1-8b-instant",
        temperature=0
    )

    response = llm.invoke(prompt_text)

    # Proper extraction
    try:
        return response.content.strip()
    except:
        return str(response).strip()

def _is_poor_title(title: str) -> bool:
    if title is None:
        return True
    cleaned = title.strip()
    if not cleaned:
        return True
    lowered = cleaned.lower()
    if lowered in {"untitled", "[image only page]"}:
        return True
    if len(cleaned) < 4:
        return True
    if sum(ch.isalpha() for ch in cleaned) < 2:
        return True
    return False

def _generate_topic_title_for_page(page_index: int, title: str, content: str) -> str:
    prompt = f"""
You are generating a concise topic title for a single PDF page.

Page index: {page_index}
Extracted title: {title or ""}
Page content:
{content or ""}

Return ONLY a concise topic title (3-8 words).
Do NOT add quotes, markdown, or extra text.
"""
    raw = run_llm_direct(prompt)
    cleaned = (raw or "").strip().strip('"').strip("'")
    if not cleaned:
        return f"Topic {page_index + 1}"
    return cleaned

def _is_bad_topic_label(text: str) -> bool:
    if text is None:
        return True
    cleaned = text.strip()
    if not cleaned:
        return True
    lowered = cleaned.lower().replace(" ", "")
    banned_fragments = [
        "notfordistribution",
        "confidential",
        "internalonly",
        "proprietary",
        "do not distribute",
        "donotdistribute",
        "copyright",
        "all rights reserved"
    ]
    return any(fragment.replace(" ", "") in lowered for fragment in banned_fragments)

def _safe_text(value: str, fallback: str) -> str:
    if not isinstance(value, str):
        return fallback
    cleaned = value.strip()
    return cleaned if cleaned else fallback

def _normalize_stage_a_output(structured: dict, slide_title: str, slide_content: str, idx_topic: int) -> dict:
    fallback_title = slide_title if not _is_poor_title(slide_title) else f"Page {idx_topic + 1}"
    topic = _safe_text(structured.get("topic") if isinstance(structured, dict) else None, fallback_title)
    if _is_bad_topic_label(topic) or _is_poor_title(topic):
        topic = fallback_title

    canvas = structured.get("canvas") if isinstance(structured, dict) else None
    if not isinstance(canvas, dict):
        canvas = {}
    canvas_title = _safe_text(canvas.get("title"), fallback_title)
    if _is_bad_topic_label(canvas_title) or _is_poor_title(canvas_title):
        canvas_title = fallback_title

    content_fallback = "This page explains a key concept from the document."
    canvas_content = _safe_text(canvas.get("content"), content_fallback)

    important_points = canvas.get("important_points") if isinstance(canvas, dict) else None
    if not isinstance(important_points, list):
        important_points = []
    important_points = [p for p in important_points if isinstance(p, str) and p.strip()]

    voice_source = structured.get("voice_source") if isinstance(structured, dict) else None
    lines = voice_source.get("lines") if isinstance(voice_source, dict) else None
    if not isinstance(lines, list):
        lines = []
    lines = [l.strip() for l in lines if isinstance(l, str) and l.strip()]
    if not lines:
        candidates = [l.strip() for l in slide_content.split("\n") if l.strip()]
        lines = candidates[:2] if candidates else [canvas_title]

    return {
        "topic": topic,
        "canvas": {
            "title": canvas_title,
            "content": canvas_content,
            "important_points": important_points
        },
        "voice_source": {
            "lines": lines
        }
    }

def _clean_professor_voice_script(script: str, max_sentences: int = 6) -> str:
    if not isinstance(script, str):
        return ""

    text = script.strip()
    if not text:
        return ""

    # Remove common filler starts and normalize whitespace.
    text = re.sub(r"^(arey|arre|dekho|chalo|toh|acha|hmm)[\s,.:;-]*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s+", " ", text).strip()

    # Keep narration concise by limiting sentence count.
    parts = re.split(r"(?<=[.!?])\s+", text)
    parts = [p.strip() for p in parts if p.strip()]
    if len(parts) > max_sentences:
        text = " ".join(parts[:max_sentences]).strip()

    # Reintroduce light pause points for TTS flow.
    text = re.sub(r"\s*\.\s*", ".\n", text)
    text = re.sub(r"\n{2,}", "\n", text).strip()
    return text

import shutil
import json
import uuid
import docx
from datetime import datetime

import uuid
from typing import List, Annotated, Optional
from fastapi import FastAPI, UploadFile, File, HTTPException, Depends
from pydantic import BaseModel
from dotenv import load_dotenv
from elevenlabs.client import ElevenLabs
from fastapi.responses import StreamingResponse
from gtts import gTTS
from io import BytesIO
from fastapi.middleware.cors import CORSMiddleware

from database import (
    init_db, get_sessions_from_db, get_sessions_with_created_at, add_session_to_db, delete_session_from_db,
    add_history_to_db, get_history_from_db, add_file_to_db, update_session_name, get_session_name, get_files_from_db,
    add_chunk_to_db, get_chunks_for_file, get_chunks_by_range, add_topic_to_db, get_topics_for_file,
    add_slide_to_db, get_slides_for_file
)


# Import auth router and user dependency
from auth import router as auth_router, get_current_user

from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_chroma import Chroma
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document

# -------------------- ENV --------------------
load_dotenv()
api_key = os.getenv("GROQ_API_KEY")
elevenlabs_api_key = os.getenv("ELEVENLABS_API_KEY")

# -------------------- TTS CLIENT --------------------
elevenlabs_client = None
if elevenlabs_api_key:
    elevenlabs_client = ElevenLabs(api_key=elevenlabs_api_key)

# -------------------- TTS FUNCTION --------------------
def text_to_speech_stream(text: str):
    """Convert text to speech and return audio stream"""
    # ElevenLabs (paid) - keep for quick switch back
    # if not elevenlabs_client:
    #     return None
    #
    # try:
    #     audio = elevenlabs_client.text_to_speech.convert(
    #         text=text,
    #         voice_id="FE4QURxZUK1rVrVK3PlK",  # Your voice ID
    #         model_id="eleven_v3",
    #         output_format="mp3_44100_128",
    #     )
    #     return audio
    # except Exception as e:
    #     print(f"TTS Error: {e}")
    #     return None

    # gTTS (free)
    try:
        tts = gTTS(text=text, lang="en", slow=False)

        buffer = BytesIO()
        tts.write_to_fp(buffer)
        buffer.seek(0)

        return buffer

    except Exception as e:
        print("TTS Error:", e)
        return None

# -------------------- EMBEDDINGS --------------------
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# -------------------- PATHS --------------------
VECTORSTORE_ROOT = "chroma_sessions"
UPLOAD_DIR = "uploaded_files"

os.makedirs(VECTORSTORE_ROOT, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# -------------------- DB --------------------
db_conn = init_db()

# -------------------- MEMORY --------------------
store = {}
vectorstores = {}
teaching_sessions = {}

def get_session_history(session: str) -> BaseChatMessageHistory:
    if session not in store:
        store[session] = ChatMessageHistory()
    return store[session]

def vectorstore_dir_for_session(session_id: str) -> str:
    return os.path.join(VECTORSTORE_ROOT, session_id)

def load_vectorstore(session_id: str):
    directory = vectorstore_dir_for_session(session_id)
    if os.path.isdir(directory):
        try:
            # Try the old method first (for existing vectorstores)
            return Chroma(persist_directory=directory, embedding_function=embeddings)
        except Exception as e:
            print(f"Error loading vectorstore for {session_id}: {e}")
            return None
    return None

def save_vectorstore(session_id: str, vectorstore_obj):
    if hasattr(vectorstore_obj, "persist"):
        vectorstore_obj.persist()
    vectorstores[session_id] = vectorstore_obj



def run_professor_prompt(session_id: str, prompt_text: str):
    chain = get_rag_chain_for_session(session_id)
    if chain is None:
        raise HTTPException(400, "Upload documents first")

    response = chain.invoke(
        {"input": prompt_text},
        config={"configurable": {"session_id": session_id}}
    )

    return response.get("answer", "").strip()


# -------------------- DOCX LOADER --------------------
def load_docx_file(file_path):
    doc = docx.Document(file_path)
    full_text = []

    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    content = "\n".join(full_text)
    return [Document(page_content=content, metadata={"source": file_path})]

# -------------------- FASTAPI --------------------

app = FastAPI(title="Graspd API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Register auth router
app.include_router(auth_router)

# -------------------- MODELS --------------------
class SessionCreate(BaseModel):
    name: str

class SessionUpdate(BaseModel):
    name: str

class ChatQuery(BaseModel):
    query: str

class TTSRequest(BaseModel):
    text: str

class MessageHistory(BaseModel):
    role: str
    content: str
    timestamp: str

class SessionInfo(BaseModel):
    session_id: str
    name: str
    created_at: Optional[str]


# -------------------- STARTUP --------------------
# (Removed default session creation; sessions are now user-specific)


# -------------------- SESSION APIs --------------------
@app.get("/sessions", response_model=List[SessionInfo])
def get_sessions(current_user: dict = Depends(get_current_user)):
    sessions = []
    for row in get_sessions_with_created_at(user_id=current_user["id"]):
        session_id = row["session_id"]
        created_at = row["created_at"]
        name = get_session_name(session_id, current_user["id"])
        sessions.append({
            "session_id": session_id,
            "name": name,
            "created_at": created_at,
        })
    return sessions


@app.post("/sessions")
def create_session(session: SessionCreate, current_user: dict = Depends(get_current_user)):
    session_id = str(uuid.uuid4())
    name = session.name.strip()

    if not name:
        raise HTTPException(400, "Session name cannot be empty")

    # No need to check for duplicate session_id, always unique
    created_at = datetime.utcnow().isoformat()
    add_session_to_db(session_id, current_user["id"], name, created_at)
    get_session_history(session_id)

    return {
        "message": "Session created",
        "session_id": session_id,
        "name": name,
        "created_at": created_at,
    }

@app.put("/sessions/{session_id}")
def update_session(session_id: str, session: SessionUpdate, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")
    
    new_name = session.name.strip()
    
    if not new_name:
        raise HTTPException(400, "Session name cannot be empty")
    
    update_session_name(session_id, new_name, current_user["id"])
    return {"message": "Session updated", "session_id": session_id, "name": new_name}

@app.delete("/sessions/{session_id}")
def delete_session(session_id: str, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")

    store.pop(session_id, None)
    vectorstores.pop(session_id, None)

    delete_session_from_db(session_id, user_id=current_user["id"])

    vector_dir = vectorstore_dir_for_session(session_id)
    if os.path.isdir(vector_dir):
        shutil.rmtree(vector_dir)

    return {"message": f"{session_id} deleted"}

# -------------------- UPLOAD API --------------------


@app.post("/sessions/{session_id}/upload")
async def upload_documents(
    session_id: str,
    files: Annotated[List[UploadFile], File(..., description="Upload PDF files")],
    current_user: dict = Depends(get_current_user)
):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(status_code=404, detail="Session not found")

    failed_files = []
    session_upload_dir = os.path.join(UPLOAD_DIR, session_id)
    os.makedirs(session_upload_dir, exist_ok=True)

    all_splits = []  # For Chroma (chat only)
    slide_db_count = 0

    for uploaded_file in files:
        file_ext = uploaded_file.filename.split('.')[-1].lower()

        # Restrict file types (MVP: PDF only)
        if uploaded_file.content_type != "application/pdf" or file_ext != "pdf":
            failed_files.append({
                "filename": uploaded_file.filename,
                "error": "Only PDF files are supported in MVP mode"
            })
            continue

        safe_filename = f"{uuid.uuid4().hex}_{uploaded_file.filename}"
        file_path = os.path.join(session_upload_dir, safe_filename)

        # Save file
        with open(file_path, "wb") as f:
            f.write(await uploaded_file.read())

        try:
            loader = PyPDFLoader(file_path)
            docs = loader.load()

            add_file_to_db(session_id, uploaded_file.filename, file_path, user_id=current_user["id"])

            # --- SLIDE/PAGE extraction for teaching (with title/body split) ---
            slides = []  # Each is a dict: {"title": ..., "content": ...}
            for idx, doc in enumerate(docs):
                content = doc.page_content.strip()
                lines = [l.strip() for l in content.split("\n") if l.strip()]
                title = lines[0] if lines else "Untitled"
                body = "\n".join(lines[1:]) if len(lines) > 1 else ""
                if not title and not body:
                    title = "[IMAGE ONLY PAGE]"
                    body = ""
                add_slide_to_db(session_id, file_path, idx, body, title, user_id=current_user["id"])
                slides.append({"title": title, "content": body})
                slide_db_count += 1

            total_slides = len(slides)
            if total_slides < 1:
                failed_files.append({
                    "filename": uploaded_file.filename,
                    "error": "No slides/pages extracted"
                })
                raise HTTPException(500, "No slides/pages extracted")

            # LEGACY TOPIC PIPELINE (POST-MVP)
            # --- Topic generation: one topic per page ---
            # for idx, slide in enumerate(slides):
            #     title = slide.get("title", "")
            #     body = slide.get("content", "")
            #     if _is_poor_title(title):
            #         if body or title:
            #             title = _generate_topic_title_for_page(idx, title, body)
            #         else:
            #             title = f"Topic {idx + 1}"
            #
            #     add_topic_to_db(
            #         session_id,
            #         file_path,
            #         idx,
            #         title,
            #         idx,
            #         idx,
            #         user_id=current_user["id"]
            #     )

            # --- Chroma chunking for chat endpoint only ---
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=2000,
                chunk_overlap=200
            )
            splits = splitter.split_documents(docs)
            all_splits.extend(splits)

        except Exception as e:
            failed_files.append({
                "filename": uploaded_file.filename,
                "error": str(e)
            })

    if not all_splits:
        raise HTTPException(status_code=400, detail={"errors": failed_files})

    # Chroma vectorstore for chat endpoint only
    vectorstore_dir = vectorstore_dir_for_session(session_id)
    os.makedirs(vectorstore_dir, exist_ok=True)
    vectorstore = Chroma.from_documents(
        documents=all_splits,
        embedding=embeddings,
        persist_directory=vectorstore_dir
    )
    save_vectorstore(session_id, vectorstore)

    return {
        "message": "Upload successful",
        "slides": slide_db_count,
        "failed_files": failed_files
    }



# --- New teaching step logic: sequential chunk fetching, no similarity_search ---
def _make_teaching_step(session_id: str, user_id: int):
    if session_id not in teaching_sessions:
        raise HTTPException(400, "Teaching session not started")

    state = teaching_sessions[session_id]
    pdfs = state.get("pdfs", [])

    if not pdfs:
        raise HTTPException(400, "No PDFs in teaching state")

    idx_pdf = state["current_pdf_index"]
    idx_topic = state["current_topic_index"]

    if idx_pdf >= len(pdfs):
        return {"message": "Teaching completed"}

    current_pdf = pdfs[idx_pdf]
    file_path = current_pdf.get("file_path")

    # MVP FLOW
    slides = get_slides_for_file(session_id, file_path, user_id)
    if not slides:
        raise HTTPException(400, "No slides found. Upload may have failed.")

    if idx_topic >= len(slides):
        return {"message": "Teaching completed"}

    current_slide = slides[idx_topic]
    slide_title = current_slide.get("title", "")
    slide_content = current_slide.get("content", "")

    # IMAGE PAGE: if page body is missing/marked image-only, do not run LLM explanation.
    is_image_only = (not slide_content or slide_content == "[IMAGE ONLY PAGE]")

    if is_image_only:
        page_title = slide_title if not _is_poor_title(slide_title) else f"Page {idx_topic + 1}"
        return {
            "step": idx_topic + 1,
            "pdf_index": idx_pdf,
            "topic_index": idx_topic,
            "topic": page_title,
            "canvas": {
                "title": page_title,
                "content": "Image page.",
                "important_points": []
            },
            "voice": {
                "script": "This is an image page."
            }
        }

    context_fragment = f"Title: {slide_title}\nContent:\n{slide_content}"
    context_fragment = context_fragment[:3500]

    # -------------------- STAGE A --------------------
    stage_a_prompt = f"""
Return ONLY valid JSON using this exact schema:

{{
  "topic": "string",
  "canvas": {{
    "title": "string",
    "content": "short explanation in English (max 5 lines)",
    "important_points": ["string", "string", "string", "string"]
  }},
  "voice_source": {{
    "lines": ["string", "string", "string", "string"]
  }}
}}

Rules:
- Valid JSON only
- No markdown
- No extra text
- Explain only what is written in the page
- Do not add outside assumptions
- If page contains multiple headings/sections, include EVERY major section
- Never ignore separate feature names like Auto Payment, Current Billing, etc.
- Summarize each section in short form
- important_points = one point per major section if possible
- voice_source.lines = include all major sections
- Do not use legal/watermark text as topic/title
- Keep concise but complete

Context:
{context_fragment}
"""

    stage_a_raw = run_llm_direct(stage_a_prompt)
    stage_a_structured = extract_json_object(stage_a_raw)

    if not stage_a_structured:
        repair_prompt = f"""
Convert this into valid JSON only using schema:

{{
  "topic": "string",
  "canvas": {{
    "title": "string",
    "content": "string",
    "important_points": ["string", "string", "string", "string"]
  }},
  "voice_source": {{
    "lines": ["string", "string", "string", "string"]
  }}
}}

Text:
{stage_a_raw}
"""
        repaired = run_llm_direct(repair_prompt)
        stage_a_structured = extract_json_object(repaired)

    if not stage_a_structured:
        stage_a_structured = {
            "topic": slide_title or f"Page {idx_topic + 1}",
            "canvas": {
                "title": slide_title or f"Page {idx_topic + 1}",
                "content": "This page explains the given topic.",
                "important_points": []
            },
            "voice_source": {
                "lines": [slide_content[:120] if slide_content else slide_title]
            }
        }

    stage_a = _normalize_stage_a_output(
        stage_a_structured,
        slide_title,
        slide_content,
        idx_topic
    )

    # -------------------- STAGE B --------------------
    source_lines = "\n".join(stage_a["voice_source"]["lines"])

    stage_b_prompt = f"""
You are an excellent professor giving a memorable short explanation.

Goal:
Teach clearly and naturally from the key concepts only.

Rules:
- Plain text only
- 4-6 concise sentences
- Natural bilingual flow (English + Roman Hindi)
- Keep technical terms in English
- Use Hindi only to improve clarity
- No robotic line-by-line translation
- No repetition or overexplaining
- Explain every concept provided. Do not skip any section.
- No filler words: arey, arre, dekho, chalo, toh, acha, hmm
- Warm, smart, student-friendly professional tone

Key concepts:
{source_lines}
"""

    voice_script = run_llm_direct(stage_b_prompt).strip()

    if not voice_script:
        voice_script = "Yeh page given concept ko short mein explain karta hai."

    voice_script = _clean_professor_voice_script(voice_script)
    if not voice_script:
        voice_script = "This page explains an important concept. Matlab is concept ko samajhna practical decision making ke liye zaroori hai."

    topic = stage_a["topic"]

    if _is_bad_topic_label(topic) or _is_poor_title(topic):
        topic = slide_title if not _is_poor_title(slide_title) else f"Page {idx_topic + 1}"

    canvas_title = stage_a["canvas"].get("title", topic)
    if _is_bad_topic_label(canvas_title) or _is_poor_title(canvas_title):
        canvas_title = topic

    stage_a["canvas"]["title"] = canvas_title

    return {
        "step": idx_topic + 1,
        "pdf_index": idx_pdf,
        "topic_index": idx_topic,
        "topic": topic,
        "canvas": stage_a["canvas"],
        "voice": {
            "script": voice_script
        }
    }
# --- New teaching start: fetch topics from DB, no similarity_search ---
@app.post("/sessions/{session_id}/teach/start")
def start_teaching(session_id: str, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")

    files = get_files_from_db(session_id, user_id=current_user["id"])
    if not files:
        raise HTTPException(400, "No PDFs uploaded for this session")

    teaching_state_pdfs = []
    for file in files:
        file_name = file["file_name"]
        file_path = file["local_path"]
        # topics will be fetched from DB in _make_teaching_step
        teaching_state_pdfs.append({"file_name": file_name, "file_path": file_path})

    teaching_sessions[session_id] = {
        "pdfs": teaching_state_pdfs,
        "current_pdf_index": 0,
        "current_topic_index": 0
    }
    return _make_teaching_step(session_id, current_user["id"])



# --- New teaching next: use DB topic count, no similarity_search ---
@app.post("/sessions/{session_id}/teach/next")
def next_teaching_step(session_id: str, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")
    if session_id not in teaching_sessions:
        raise HTTPException(400, "Teaching not started")

    state = teaching_sessions[session_id]
    state["current_topic_index"] += 1

    # MVP FLOW: advance by page count
    if state["current_pdf_index"] < len(state["pdfs"]):
        current_pdf = state["pdfs"][state["current_pdf_index"]]
        file_path = current_pdf["file_path"]
        slides = get_slides_for_file(session_id, file_path, current_user["id"])
        if state["current_topic_index"] >= len(slides):
            state["current_pdf_index"] += 1
            state["current_topic_index"] = 0

    # LEGACY TOPIC PIPELINE (POST-MVP)
    # if state["current_pdf_index"] < len(state["pdfs"]):
    #     current_pdf = state["pdfs"][state["current_pdf_index"]]
    #     file_path = current_pdf["file_path"]
    #     topics_db = get_topics_for_file(session_id, file_path, current_user["id"])
    #     if state["current_topic_index"] >= len(topics_db):
    #         state["current_pdf_index"] += 1
    #         state["current_topic_index"] = 0

    if state["current_pdf_index"] >= len(state["pdfs"]):
        return {"message": "Teaching completed"}

    teaching_sessions[session_id] = state
    return _make_teaching_step(session_id, current_user["id"])


# -------------------- TTS API --------------------
@app.post("/tts")
async def generate_tts(tts_request: TTSRequest):
    """Convert text to speech using gTTS (fallback) or ElevenLabs"""
    
    text = tts_request.text.strip()
    if not text:
        raise HTTPException(400, "Text cannot be empty")
    
    audio_stream = text_to_speech_stream(text)
    if audio_stream is None:
        raise HTTPException(500, "Failed to generate audio")
    
    return StreamingResponse(audio_stream, media_type="audio/mpeg")


# -------------------- RAG --------------------
def get_rag_chain_for_session(session_id: str):
    if not api_key:
        raise Exception("Missing GROQ_API_KEY")

    llm = ChatGroq(
        groq_api_key=api_key,
        model_name="llama-3.1-8b-instant"
    )

    vs = vectorstores.get(session_id) or load_vectorstore(session_id)

    if vs is None:
        return None

    vectorstores[session_id] = vs
    retriever = vs.as_retriever()

    contextualize_prompt = ChatPromptTemplate.from_messages([
        ("system", "Rephrase question with history"),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}")
    ])

    history_aware_retriever = create_history_aware_retriever(
        llm, retriever, contextualize_prompt
    )

    qa_prompt = ChatPromptTemplate.from_messages([
        ("system", "Answer using context. Max 3 sentences.\n{context}"),
        MessagesPlaceholder("chat_history"),
        ("human", "{input}")
    ])

    qa_chain = create_stuff_documents_chain(llm, qa_prompt)

    rag_chain = create_retrieval_chain(
        history_aware_retriever, qa_chain
    )

    return RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key="input",
        history_messages_key="chat_history",
        output_messages_key="answer"
    )

# -------------------- CHAT --------------------
@app.post("/sessions/{session_id}/chat")
def chat(session_id: str, query: ChatQuery, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")

    chain = get_rag_chain_for_session(session_id)

    if chain is None:
        raise HTTPException(400, "Upload documents first")

    user_input = query.query.strip()
    if not user_input:
        raise HTTPException(400, "Empty query")

    add_history_to_db(session_id, "user", user_input, user_id=current_user["id"])

    response = chain.invoke(
        {"input": user_input},
        config={"configurable": {"session_id": session_id}}
    )

    answer = response.get("answer", "")

    add_history_to_db(session_id, "assistant", answer, user_id=current_user["id"])

    return {"user": user_input, "assistant": answer}

# -------------------- HISTORY --------------------
@app.get("/sessions/{session_id}/history", response_model=List[MessageHistory])
def get_history(session_id: str, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")

    rows = get_history_from_db(session_id, user_id=current_user["id"])

    return [
        {
            "role": r["role"],
            "content": r["content"],
            "timestamp": r["timestamp"]
        }
        for r in rows
    ]

# -------------------- DOCUMENTS --------------------
@app.get("/sessions/{session_id}/documents")
def get_documents(session_id: str, current_user: dict = Depends(get_current_user)):
    if session_id not in get_sessions_from_db(user_id=current_user["id"]):
        raise HTTPException(404, "Session not found")

    rows = get_files_from_db(session_id, user_id=current_user["id"])

    documents = [
        {
            "file_name": r["file_name"],
            "local_path": r["local_path"],
            "uploaded_at": r["uploaded_at"]
        }
        for r in rows
    ]

    return {"documents": documents}