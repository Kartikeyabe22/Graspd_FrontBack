import streamlit as st
import os
from dotenv import load_dotenv

# Local imports
from database import init_db, get_sessions_from_db, add_session_to_db, delete_session_from_db, add_history_to_db, get_history_from_db, add_file_to_db

# LangChain imports
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_chroma import Chroma
from langchain_community.chat_message_histories import ChatMessageHistory
from langchain_core.chat_history import BaseChatMessageHistory
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
from langchain_groq import ChatGroq
from langchain_core.runnables.history import RunnableWithMessageHistory
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
import docx

# Custom DOCX loader function
def load_docx_file(file_path):
    """Load a DOCX file and return a Document object."""
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    
    content = '\n'.join(full_text)
    return [Document(page_content=content, metadata={"source": file_path})]

# Load env
load_dotenv()

# Set HF token
os.environ['HF_TOKEN'] = os.getenv("HF_TOKEN")

# Embeddings
embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")

# UI
st.title("Graspd")

# Load API key
api_key = os.getenv("GROQ_API_KEY")

# persistence config
VECTORSTORE_ROOT = "chroma_sessions"

os.makedirs(VECTORSTORE_ROOT, exist_ok=True)

# Initialize database
init_db()


# Initialize session stores
if 'store' not in st.session_state:
    st.session_state.store = {}

if 'vectorstores' not in st.session_state:
    st.session_state.vectorstores = {}

existing_sessions = get_sessions_from_db()
if existing_sessions:
    st.session_state.sessions = existing_sessions
else:
    st.session_state.sessions = ["default_session"]
    add_session_to_db("default_session")

if 'current_session' not in st.session_state:
    st.session_state.current_session = st.session_state.sessions[0]


# ---------------- SESSION FUNCTION ---------------- #
def get_session_history(session: str) -> BaseChatMessageHistory:
    if 'store' not in st.session_state:
        st.session_state.store = {}

    if session not in st.session_state.store:
        st.session_state.store[session] = ChatMessageHistory()

    return st.session_state.store[session]


def vectorstore_dir_for_session(session_id: str) -> str:
    return os.path.join(VECTORSTORE_ROOT, session_id)


def load_vectorstore(session_id: str):
    directory = vectorstore_dir_for_session(session_id)
    if os.path.isdir(directory):
        try:
            vs = Chroma(persist_directory=directory, embedding_function=embeddings)
            return vs
        except Exception:
            return None
    return None


def save_vectorstore(session_id: str, vectorstore_obj):
    if hasattr(vectorstore_obj, 'persist'):
        vectorstore_obj.persist()
    st.session_state.vectorstores[session_id] = vectorstore_obj


# ---------------- MAIN APP ---------------- #
if api_key:

    llm = ChatGroq(
        groq_api_key=api_key,
        model_name="llama-3.3-70b-versatile"
    )

    with st.sidebar:
        st.header("Chat sessions")

        new_session = st.text_input("New session name", key="new_session_name")
        if st.button("Create session", key="create_session_btn"):
            new_session = new_session.strip()
            if new_session:
                if new_session not in st.session_state.sessions:
                    st.session_state.sessions.append(new_session)
                    st.session_state.current_session = new_session
                    add_session_to_db(new_session)
                    get_session_history(new_session)
                    st.success(f"Created session '{new_session}'")
                    # The new session name is kept in the input widget value automatically;
                    # avoid direct write to st.session_state from inside callback.
                else:
                    st.warning("Session already exists")

        if st.session_state.current_session not in st.session_state.sessions:
            st.session_state.current_session = st.session_state.sessions[0]

        session_index = st.session_state.sessions.index(st.session_state.current_session)
        selected_session = st.radio("Select session", st.session_state.sessions, index=session_index, key="session_picker")
        st.session_state.current_session = selected_session

        if len(st.session_state.sessions) > 1 and st.button("Delete current session", key="delete_session_btn"):
            delete_name = st.session_state.current_session
            if delete_name in st.session_state.sessions:
                st.session_state.sessions.remove(delete_name)
                st.session_state.store.pop(delete_name, None)
                st.session_state.vectorstores.pop(delete_name, None)
                delete_session_from_db(delete_name)
                vector_dir = vectorstore_dir_for_session(delete_name)
                if os.path.isdir(vector_dir):
                    try:
                        import shutil
                        shutil.rmtree(vector_dir)
                    except Exception:
                        pass
                st.session_state.current_session = st.session_state.sessions[0]
                st.success(f"Deleted session '{delete_name}'")

        st.markdown("### Stored sessions")
        for session_name in st.session_state.sessions:
            st.write(f"• {session_name}")

    session_id = st.session_state.current_session
    st.subheader(f"Active session: {session_id}")

    uploaded_files = st.file_uploader(
        "Upload PDF or DOCX files",
        type=["pdf", "doc", "docx"],
        accept_multiple_files=True,
        key=f"file_uploader_{session_id}"
    )

    conversational_rag_chain = None  # safety

    # Load vectorstore for selected session if exists
    session_vectorstore = st.session_state.vectorstores.get(session_id) or load_vectorstore(session_id)
    if session_vectorstore is not None:
        st.session_state.vectorstores[session_id] = session_vectorstore
        retriever = session_vectorstore.as_retriever()
        contextualize_q_prompt = ChatPromptTemplate.from_messages([
            ("system", "Given chat history and latest question, reformulate into standalone question."),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}")
        ])

        history_aware_retriever = create_history_aware_retriever(
            llm,
            retriever,
            contextualize_q_prompt
        )

        qa_prompt = ChatPromptTemplate.from_messages([
            ("system", "You are an assistant. Use context to answer. If unknown, say so. Max 3 sentences.\n\n{context}"),
            MessagesPlaceholder("chat_history"),
            ("human", "{input}")
        ])

        question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)

        rag_chain = create_retrieval_chain(
            history_aware_retriever,
            question_answer_chain
        )

        conversational_rag_chain = RunnableWithMessageHistory(
            rag_chain,
            get_session_history,
            input_messages_key="input",
            history_messages_key="chat_history",
            output_messages_key="answer"
        )

    # ---------------- PROCESS DOCUMENTS ---------------- #
    if uploaded_files:
        documents = []
        
        # Ensure upload dict exists
        UPLOAD_DIR = "uploaded_files"
        session_upload_dir = os.path.join(UPLOAD_DIR, session_id)
        os.makedirs(session_upload_dir, exist_ok=True)
        
        import uuid
        for uploaded_file in uploaded_files:
            safe_filename = f"{uuid.uuid4().hex}_{uploaded_file.name}"
            file_path = os.path.join(session_upload_dir, safe_filename)

            with open(file_path, "wb") as f:
                f.write(uploaded_file.getvalue())

            # Determine file type and use appropriate loader
            file_extension = uploaded_file.name.split('.')[-1].lower()

            try:
                if file_extension == 'pdf':
                    loader = PyPDFLoader(file_path)
                    docs = loader.load()
                elif file_extension == 'docx':
                    docs = load_docx_file(file_path)
                elif file_extension == 'doc':
                    try:
                        docs = load_docx_file(file_path)
                    except Exception:
                        st.error(f"Cannot process '{uploaded_file.name}'. The older .doc format is not supported locally. Please save it as .docx first.")
                        if os.path.exists(file_path):
                            os.remove(file_path)
                        continue
                else:
                    st.error(f"Unsupported file type: {file_extension}. Please upload PDF or DOCX files.")
                    if os.path.exists(file_path):
                        os.remove(file_path)
                    continue

                documents.extend(docs)

                # save upload metadata
                add_file_to_db(session_id, uploaded_file.name, file_path)

            except Exception as e:
                st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
                if os.path.exists(file_path):
                    os.remove(file_path)
                continue

        # Check if any documents were successfully loaded
        if not documents:
            st.error("No valid documents could be extracted. Please check your files and try again.")

        else:
            # Split
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=5000,
                chunk_overlap=500
            )
            splits = text_splitter.split_documents(documents)

            # Vector DB persisted per session
            vectorstore_dir = vectorstore_dir_for_session(session_id)
            os.makedirs(vectorstore_dir, exist_ok=True)

            vectorstore = Chroma.from_documents(
                documents=splits,
                embedding=embeddings,
                persist_directory=vectorstore_dir
            )

            save_vectorstore(session_id, vectorstore)

            retriever = vectorstore.as_retriever()

            contextualize_q_prompt = ChatPromptTemplate.from_messages([
                ("system", "Given chat history and latest question, reformulate into standalone question."),
                MessagesPlaceholder("chat_history"),
                ("human", "{input}")
            ])

            history_aware_retriever = create_history_aware_retriever(
                llm,
                retriever,
                contextualize_q_prompt
            )

            qa_prompt = ChatPromptTemplate.from_messages([
                ("system", "You are an assistant. Use context to answer. If unknown, say so. Max 3 sentences.\n\n{context}"),
                MessagesPlaceholder("chat_history"),
                ("human", "{input}")
            ])

            question_answer_chain = create_stuff_documents_chain(llm, qa_prompt)

            rag_chain = create_retrieval_chain(
                history_aware_retriever,
                question_answer_chain
            )

            conversational_rag_chain = RunnableWithMessageHistory(
                rag_chain,
                get_session_history,
                input_messages_key="input",
                history_messages_key="chat_history",
                output_messages_key="answer"
            )

    # ---------------- CHAT ---------------- #
    history_rows = get_history_from_db(session_id)
    # if history_rows:
    #     st.markdown("### Conversation history")
    #     for row in history_rows:
    #         st.markdown(f"**{row['role'].capitalize()}:** {row['content']}")

    user_input = st.text_input("Ask something from the PDF:")

    if user_input:
        if conversational_rag_chain is None:
            st.warning("⚠️ Please upload PDF first")
        else:
            add_history_to_db(session_id, "user", user_input, user_id=None)
            response = conversational_rag_chain.invoke(
                {"input": user_input},
                config={"configurable": {"session_id": session_id}}
            )
            answer_text = response.get("answer", "")
            add_history_to_db(session_id, "assistant", answer_text, user_id=None)

            st.write("### 🧑 You:")
            st.write(user_input)

            st.write("### 🤖 Assistant:")
            st.write(answer_text)

else:
    st.warning("Please set GROQ_API_KEY in .env")